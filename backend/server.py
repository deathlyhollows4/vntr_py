"""SmartLegal-AI Backend — Contract & Legal Document RAG Analyzer."""
from __future__ import annotations

import io
import json
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from dotenv import load_dotenv
from emergentintegrations.llm.chat import LlmChat, UserMessage
from fastapi import APIRouter, FastAPI, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field
from pypdf import PdfReader
from rank_bm25 import BM25Okapi
from starlette.middleware.cors import CORSMiddleware

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger("smartlegal")

MONGO_URL = os.environ["MONGO_URL"]
DB_NAME = os.environ["DB_NAME"]
EMERGENT_LLM_KEY = os.environ["EMERGENT_LLM_KEY"]
LLM_MODEL = "gpt-4o"

client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

app = FastAPI(title="SmartLegal-AI")
api = APIRouter(prefix="/api")


# ---------- Utility helpers ----------
def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def tokenize(text: str) -> list[str]:
    return re.findall(r"[A-Za-z0-9']+", text.lower())


def chunk_pages(pages: list[str], size: int = 900, overlap: int = 150) -> list[dict]:
    """Chunk per-page text into overlapping passages and keep page provenance."""
    chunks: list[dict] = []
    for page_idx, text in enumerate(pages, start=1):
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(start + size, len(text))
            if end < len(text):
                # prefer breaking at a sentence boundary
                dot = text.rfind(". ", start, end)
                if dot > start + size // 2:
                    end = dot + 1
            passage = text[start:end].strip()
            if passage:
                chunks.append({"id": str(uuid.uuid4()), "page": page_idx, "text": passage})
            if end >= len(text):
                break
            start = max(end - overlap, start + 1)
    return chunks


def extract_pdf(raw: bytes) -> list[str]:
    reader = PdfReader(io.BytesIO(raw))
    out = []
    for p in reader.pages:
        try:
            out.append(p.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            logger.warning("PDF page extraction failed: %s", exc)
            out.append("")
    return out


def bm25_search(chunks: list[dict], query: str, k: int = 5) -> list[dict]:
    if not chunks:
        return []
    corpus = [tokenize(c["text"]) for c in chunks]
    bm25 = BM25Okapi(corpus)
    scores = bm25.get_scores(tokenize(query))
    idxs = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:k]
    return [chunks[i] for i in idxs if scores[i] > 0] or [chunks[i] for i in idxs[:k]]


async def llm_complete(system: str, user: str, session_id: str) -> str:
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message=system).with_model(
        "openai", LLM_MODEL
    )
    resp = await chat.send_message(UserMessage(text=user))
    return resp if isinstance(resp, str) else str(resp)


def safe_json_loads(raw: str) -> Any:
    raw = raw.strip()
    # strip code fences
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    # find outermost JSON object/array
    for pat in (r"\{.*\}", r"\[.*\]"):
        m = re.search(pat, raw, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:  # noqa: BLE001
                continue
    try:
        return json.loads(raw)
    except Exception:  # noqa: BLE001
        return None


# ---------- Pydantic Schemas ----------
class QueryIn(BaseModel):
    question: str


class CompareIn(BaseModel):
    doc1_id: str
    doc2_id: str
    aspect: str = Field(default="overall terms, risks, and key differences")


class DocumentOut(BaseModel):
    id: str
    name: str
    pages: int
    chars: int
    uploaded_at: str
    status: str


# ---------- Endpoints ----------
@api.get("/")
async def root():
    return {"app": "SmartLegal-AI", "status": "ok"}


@api.post("/documents/upload", response_model=DocumentOut)
async def upload_document(file: UploadFile = File(...)):
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Only PDF files supported")
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "Empty file")
    pages = extract_pdf(raw)
    if not any(p.strip() for p in pages):
        raise HTTPException(400, "Could not extract any text from this PDF (is it scanned?)")
    chunks = chunk_pages(pages)
    doc_id = str(uuid.uuid4())
    char_count = sum(len(p) for p in pages)
    doc = {
        "id": doc_id,
        "name": file.filename,
        "pages_count": len(pages),
        "chars": char_count,
        "uploaded_at": now_iso(),
        "status": "ready",
        "chunks": chunks,
        "page_texts": pages,
    }
    await db.documents.insert_one(doc)
    logger.info("Uploaded '%s' id=%s pages=%d chunks=%d", file.filename, doc_id, len(pages), len(chunks))
    return DocumentOut(
        id=doc_id, name=file.filename, pages=len(pages), chars=char_count, uploaded_at=doc["uploaded_at"], status="ready"
    )


@api.get("/documents", response_model=list[DocumentOut])
async def list_documents():
    docs = await db.documents.find(
        {}, {"_id": 0, "id": 1, "name": 1, "pages_count": 1, "chars": 1, "uploaded_at": 1, "status": 1}
    ).sort("uploaded_at", -1).to_list(500)
    return [
        DocumentOut(
            id=d["id"], name=d["name"], pages=d.get("pages_count", 0), chars=d.get("chars", 0),
            uploaded_at=d.get("uploaded_at", ""), status=d.get("status", "ready"),
        )
        for d in docs
    ]


async def _load_doc(doc_id: str) -> dict:
    doc = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Document not found")
    return doc


@api.get("/documents/{doc_id}")
async def get_document(doc_id: str):
    doc = await _load_doc(doc_id)
    return {
        "id": doc["id"],
        "name": doc["name"],
        "pages": doc.get("pages_count", 0),
        "chars": doc.get("chars", 0),
        "uploaded_at": doc.get("uploaded_at", ""),
        "status": doc.get("status", "ready"),
    }


@api.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    await db.documents.delete_one({"id": doc_id})
    await db.chat_history.delete_many({"document_id": doc_id})
    await db.analyses.delete_many({"document_id": doc_id})
    return {"deleted": doc_id}


@api.post("/documents/{doc_id}/query")
async def query_document(doc_id: str, body: QueryIn):
    doc = await _load_doc(doc_id)
    if not body.question.strip():
        raise HTTPException(400, "Empty question")

    top = bm25_search(doc.get("chunks", []), body.question, k=5)
    context = "\n\n".join(f"[Page {c['page']}]\n{c['text']}" for c in top)

    system = (
        "You are SmartLegal-AI, a careful contract and legal document analysis assistant. "
        "Answer ONLY using the provided context excerpts. If the context is insufficient, say so honestly. "
        "Always cite page numbers inline using the format [Page N] for every factual claim. "
        "Be concise, precise, and use plain legal English."
    )
    user = f"CONTEXT FROM DOCUMENT '{doc['name']}':\n{context}\n\nQUESTION: {body.question}\n\nProvide a clear answer with inline [Page N] citations."
    answer = await llm_complete(system, user, session_id=f"qa-{doc_id}")

    citations = [{"page": c["page"], "quote": c["text"][:220]} for c in top]
    record = {
        "id": str(uuid.uuid4()),
        "document_id": doc_id,
        "question": body.question,
        "answer": answer,
        "citations": citations,
        "timestamp": now_iso(),
    }
    await db.chat_history.insert_one(record.copy())
    record.pop("_id", None)
    return record


@api.get("/documents/{doc_id}/history")
async def get_history(doc_id: str):
    rows = await db.chat_history.find({"document_id": doc_id}, {"_id": 0}).sort("timestamp", 1).to_list(500)
    return rows


async def _cached_analysis(doc_id: str, kind: str, factory) -> dict:
    existing = await db.analyses.find_one({"document_id": doc_id, "type": kind}, {"_id": 0})
    if existing:
        return existing
    data = await factory()
    record = {"id": str(uuid.uuid4()), "document_id": doc_id, "type": kind, "data": data, "created_at": now_iso()}
    await db.analyses.insert_one(record.copy())
    record.pop("_id", None)
    return record


def _doc_context(doc: dict, max_chars: int = 40000) -> str:
    pieces = []
    total = 0
    for idx, text in enumerate(doc.get("page_texts", []), start=1):
        snippet = f"[Page {idx}]\n{text}\n"
        if total + len(snippet) > max_chars:
            break
        pieces.append(snippet)
        total += len(snippet)
    return "\n".join(pieces)


@api.post("/documents/{doc_id}/analyze/risks")
async def analyze_risks(doc_id: str):
    doc = await _load_doc(doc_id)

    async def factory():
        system = (
            "You are SmartLegal-AI, a senior contract risk analyst. Identify risk clauses in the provided contract. "
            "Return STRICT JSON with this shape: "
            '{"risks":[{"title":str,"severity":"High"|"Medium"|"Low","page":int,"clause":str,"why":str,"recommendation":str}]} '
            "Include between 4 and 12 risks covering liability, indemnity, termination, IP, payment, confidentiality, "
            "data/privacy, jurisdiction, auto-renewal, assignment and warranty where applicable. No prose outside JSON."
        )
        user = f"CONTRACT '{doc['name']}':\n{_doc_context(doc)}"
        raw = await llm_complete(system, user, session_id=f"risks-{doc_id}")
        parsed = safe_json_loads(raw) or {"risks": []}
        if isinstance(parsed, list):
            parsed = {"risks": parsed}
        return parsed

    return await _cached_analysis(doc_id, "risks", factory)


@api.post("/documents/{doc_id}/analyze/clauses")
async def analyze_clauses(doc_id: str):
    doc = await _load_doc(doc_id)

    async def factory():
        system = (
            "You are SmartLegal-AI. Extract key clauses from the contract grouped by category. "
            "Return STRICT JSON: "
            '{"categories":[{"name":str,"items":[{"title":str,"page":int,"text":str}]}]} '
            "Use these categories when present: Confidentiality / NDA, Payment & Fees, Intellectual Property, "
            "Termination, Liability & Indemnity, Governing Law & Jurisdiction, Warranties, Data Protection, "
            "Dispute Resolution. Keep clause text to 1-3 sentences, quoted or lightly paraphrased. No prose."
        )
        user = f"CONTRACT '{doc['name']}':\n{_doc_context(doc)}"
        raw = await llm_complete(system, user, session_id=f"clauses-{doc_id}")
        parsed = safe_json_loads(raw) or {"categories": []}
        if isinstance(parsed, list):
            parsed = {"categories": parsed}
        return parsed

    return await _cached_analysis(doc_id, "clauses", factory)


@api.post("/documents/{doc_id}/analyze/summary")
async def analyze_summary(doc_id: str):
    doc = await _load_doc(doc_id)

    async def factory():
        system = (
            "You are SmartLegal-AI, preparing a crisp executive summary for a decision-maker. "
            "Return STRICT JSON: "
            '{"overview":str,"parties":[str],"effective_date":str,"term":str,"governing_law":str,'
            '"key_obligations":[str],"financials":[str],"highlights":[str],"red_flags":[str]} '
            "Keep lists concise (3-6 bullets). Use 'Not specified' if unknown. No prose outside JSON."
        )
        user = f"CONTRACT '{doc['name']}':\n{_doc_context(doc)}"
        raw = await llm_complete(system, user, session_id=f"summary-{doc_id}")
        parsed = safe_json_loads(raw) or {"overview": raw[:2000]}
        return parsed

    return await _cached_analysis(doc_id, "summary", factory)


@api.post("/documents/compare")
async def compare_documents(body: CompareIn):
    d1 = await _load_doc(body.doc1_id)
    d2 = await _load_doc(body.doc2_id)
    system = (
        "You are SmartLegal-AI comparing two contracts side-by-side. "
        "Return STRICT JSON: "
        '{"summary":str,"rows":[{"aspect":str,"doc1":str,"doc2":str,"verdict":"doc1"|"doc2"|"equal"|"n/a"}]} '
        "Cover at least 8 aspects: parties, term & termination, payment terms, liability cap, indemnity, "
        "IP ownership, confidentiality, governing law, dispute resolution, warranties. No prose outside JSON."
    )
    user = (
        f"ASPECT FOCUS: {body.aspect}\n\n"
        f"--- DOCUMENT A: {d1['name']} ---\n{_doc_context(d1, max_chars=18000)}\n\n"
        f"--- DOCUMENT B: {d2['name']} ---\n{_doc_context(d2, max_chars=18000)}"
    )
    raw = await llm_complete(system, user, session_id=f"cmp-{body.doc1_id}-{body.doc2_id}")
    parsed = safe_json_loads(raw) or {"summary": raw[:2000], "rows": []}
    return {"doc1": {"id": d1["id"], "name": d1["name"]}, "doc2": {"id": d2["id"], "name": d2["name"]}, **parsed}


@api.get("/documents/{doc_id}/report")
async def download_report(doc_id: str):
    doc = await _load_doc(doc_id)

    summary = await db.analyses.find_one({"document_id": doc_id, "type": "summary"}, {"_id": 0})
    risks = await db.analyses.find_one({"document_id": doc_id, "type": "risks"}, {"_id": 0})
    clauses = await db.analyses.find_one({"document_id": doc_id, "type": "clauses"}, {"_id": 0})

    docx = DocxDocument()
    docx.add_heading("SmartLegal-AI — Contract Analysis Report", level=0)
    docx.add_paragraph(f"Document: {doc['name']}")
    docx.add_paragraph(f"Pages: {doc.get('pages_count', 0)} | Generated: {now_iso()}")

    if summary and summary.get("data"):
        s = summary["data"]
        docx.add_heading("Executive Summary", level=1)
        docx.add_paragraph(s.get("overview", ""))
        for key in ("parties", "key_obligations", "financials", "highlights", "red_flags"):
            if s.get(key):
                docx.add_heading(key.replace("_", " ").title(), level=2)
                for item in s[key] if isinstance(s[key], list) else [s[key]]:
                    docx.add_paragraph(str(item), style="List Bullet")
        for key in ("effective_date", "term", "governing_law"):
            if s.get(key):
                docx.add_paragraph(f"{key.replace('_',' ').title()}: {s[key]}")

    if risks and risks.get("data", {}).get("risks"):
        docx.add_heading("Risk Analysis", level=1)
        for r in risks["data"]["risks"]:
            docx.add_heading(f"[{r.get('severity','?')}] {r.get('title','Risk')} (Page {r.get('page','?')})", level=2)
            if r.get("clause"):
                docx.add_paragraph(f"Clause: {r['clause']}")
            if r.get("why"):
                docx.add_paragraph(f"Why it matters: {r['why']}")
            if r.get("recommendation"):
                docx.add_paragraph(f"Recommendation: {r['recommendation']}")

    if clauses and clauses.get("data", {}).get("categories"):
        docx.add_heading("Clause Extraction", level=1)
        for cat in clauses["data"]["categories"]:
            docx.add_heading(cat.get("name", "Category"), level=2)
            for item in cat.get("items", []):
                docx.add_paragraph(
                    f"{item.get('title','')} (Page {item.get('page','?')}): {item.get('text','')}",
                    style="List Bullet",
                )

    if not (summary or risks or clauses):
        docx.add_paragraph("Run analyses (Summary / Risks / Clauses) before downloading a full report.")

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", doc["name"]).rsplit(".", 1)[0] or "report"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="SmartLegal_{safe_name}.docx"'},
    )


app.include_router(api)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("shutdown")
async def _shutdown():
    client.close()
